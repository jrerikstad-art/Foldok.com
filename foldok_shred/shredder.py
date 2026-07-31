"""The shredder itself.

    shred = Shredder().shred("some_manual.pdf", grade="exemplary")
    print(shred.report())
    shred.proposals  # offered to the console queue, never applied

The whole flow, and the order matters:

    read bytes -> measure -> build proposals -> **drop the text** -> return

The text is a local variable in one function. It is never assigned to the
result, never cached, never written anywhere. ``Shred`` has no field that could
hold it. When the function returns, the only thing that survives is numbers,
section titles, and — for a graded reference — clause citations.

Section titles are the one judgement call, so it is worth stating: a heading like
"Verification by testing" is a name for a topic, not the expression the standards
body sells. Titles are kept because a skeleton without them is useless, and
because "documents of this kind contain a section about verification" is a fact
about convention. Bodies are not kept, ever.
"""

from __future__ import annotations

import hashlib
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable, Sequence

from .measure import design_from_docx, design_from_pdf, skeleton
from .model import (
    GRADE_LEARNS,
    DesignProfile,
    Grade,
    Shred,
    ShredProposal,
    Skeleton,
)

TEXTUAL = {".txt", ".md", ".markdown", ".rst", ".html", ".htm", ".csv", ".json"}


class ShredRefused(Exception):
    """The document cannot be measured, with the reason."""


@dataclass
class Shredder:
    min_sections: int = 3  # below this, a skeleton is not a pattern
    max_bytes: int = 80 * 1024 * 1024

    # -- the one public path ------------------------------------------------
    def shred(
        self,
        path: str | Path,
        *,
        grade: Grade = "sample",
        kind: str = "",
        standard: str = "",
    ) -> Shred:
        p = Path(path)
        if not p.exists():
            raise ShredRefused(f"{p} does not exist")
        size = p.stat().st_size
        if size > self.max_bytes:
            raise ShredRefused(
                f"{size // 1024 // 1024} MB is beyond what the shredder will read in one go"
            )

        source_id = _source_id(p)
        text = _read_text(p)  # local only, dropped on return
        design = _measure_design(p)
        bones = skeleton(text) if text else Skeleton()
        result = Shred(
            source_id=source_id, grade=grade, kind=kind or _guess_kind(bones),
            skeleton=bones, design=design, bytes_in=size,
        )

        learns = GRADE_LEARNS.get(grade, ())
        if "skeleton" in learns:
            result.proposals.extend(_skeleton_proposals(bones, source_id, self.min_sections))
        if "design" in learns:
            result.proposals.extend(_design_proposals(design, source_id))
        if "obligations" in learns:
            result.obligations = _obligations(text, standard or kind or "the standard", source_id)
            result.proposals.extend(_obligation_proposals(result.obligations, source_id))

        if not text:
            result.notes.append(
                "no text could be read — a scanned document needs OCR before it can be measured"
            )
        if grade in ("sample", "ours"):
            result.notes.append(
                f"grade '{grade}' measures but learns nothing; regrade it 'exemplary' if this "
                "is a document you want Foldok to be more like"
            )
        result.notes.append("the document itself was not retained")

        del text  # explicit, and the point of the exercise
        return result

    def shred_many(
        self, paths: Iterable[str | Path], *, grade: Grade = "sample"
    ) -> list[Shred]:
        out: list[Shred] = []
        for path in paths:
            try:
                out.append(self.shred(path, grade=grade))
            except ShredRefused:
                continue
        return out


# ----------------------------------------------------------------------
def consensus(shreds: Sequence[Shred], *, min_documents: int = 3) -> list[ShredProposal]:
    """What several exemplary documents agree on.

    One well-made manual is one person's taste. Three that agree is a
    convention, and a convention is the only thing worth turning into a default.
    """
    usable = [s for s in shreds if "skeleton" in s.learns and s.skeleton.section_count]
    if len(usable) < min_documents:
        return []

    counts: dict[str, int] = {}
    for shred in usable:
        for title in {t.strip().lower() for t in shred.skeleton.titles(level=1)}:
            counts[title] = counts.get(title, 0) + 1

    common = [(t, n) for t, n in counts.items() if n >= min_documents]
    common.sort(key=lambda kv: (-kv[1], kv[0]))
    if not common:
        return []

    designs = [s.design for s in usable if s.design.usable]
    proposals = [
        ShredProposal(
            kind="skeleton",
            title=f"{len(common)} section(s) appear in {min_documents}+ documents",
            detail=", ".join(t for t, _ in common[:8]),
            payload={"sections": [t for t, _ in common], "documents": len(usable)},
            confidence=min(0.95, 0.5 + 0.1 * len(usable)),
            source=f"{len(usable)} shreds",
        )
    ]
    if len(designs) >= min_documents:
        widths = sorted({d.page_size for d in designs})
        proposals.append(
            ShredProposal(
                kind="design",
                title="a shared page setup across the sample",
                detail=f"page {', '.join(widths)}, body "
                       f"{sum(d.body_size_pt for d in designs) / len(designs):.1f} pt",
                payload={"grid": designs[0].to_grid()},
                confidence=0.6,
                source=f"{len(designs)} shreds",
            )
        )
    return proposals


# ----------------------------------------------------------------------
def _read_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in TEXTUAL:
        return path.read_text(encoding="utf-8", errors="replace")
    if suffix == ".pdf":
        try:
            from pypdf import PdfReader
            return "\n".join((page.extract_text() or "") for page in PdfReader(str(path)).pages)
        except Exception:  # noqa: BLE001
            return ""
    if suffix == ".docx":
        try:
            import docx  # type: ignore
            document = docx.Document(str(path))
            return "\n".join(p.text for p in document.paragraphs)
        except Exception:  # noqa: BLE001
            return ""
    return ""


def _measure_design(path: Path) -> DesignProfile:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return design_from_pdf(path)
    if suffix == ".docx":
        return design_from_docx(path)
    return DesignProfile(measured_from="structure", confidence=0.0)


def _source_id(path: Path) -> str:
    """A hash of the content, so the same document shredded twice is recognised
    and the file name — often a client and a job number — is never stored."""
    h = hashlib.sha256()
    with path.open("rb") as fh:
        while True:
            block = fh.read(1 << 20)
            if not block:
                break
            h.update(block)
    return h.hexdigest()[:16]


def _guess_kind(bones: Skeleton) -> str:
    titles = " ".join(bones.titles()).lower()
    for kind, needles in (
        ("installation_manual", ("installation", "montering", "commissioning")),
        ("inspection_report", ("inspection", "inspeksjon", "findings", "avvik")),
        ("technical_file", ("risk assessment", "declaration", "conformity")),
        ("test_report", ("test", "måling", "verification", "verifikasjon")),
        ("specification", ("specification", "spesifikasjon", "requirements")),
    ):
        if any(n in titles for n in needles):
            return kind
    return ""


def _skeleton_proposals(bones: Skeleton, source: str, minimum: int) -> list[ShredProposal]:
    if bones.section_count < minimum:
        return []
    top = bones.titles(level=1)[:12]
    return [
        ShredProposal(
            kind="skeleton",
            title=f"a {bones.section_count}-section structure, {bones.depth} level(s) deep",
            detail=", ".join(top[:6]) + ("…" if len(top) > 6 else ""),
            payload={
                "sections": top,
                "depth": bones.depth,
                "numbering": bones.numbering,
                "tables": bones.tables,
                "figures": bones.figures,
            },
            confidence=0.5 if bones.section_count < 6 else 0.75,
            source=source,
        )
    ]


def _design_proposals(design: DesignProfile, source: str) -> list[ShredProposal]:
    if not design.usable:
        return []
    return [
        ShredProposal(
            kind="design",
            title=f"page setup measured from the document ({design.measured_from})",
            detail=(
                f"{design.page_size}, margins "
                f"{design.margin_left_pt:.0f}/{design.margin_top_pt:.0f} pt, "
                f"body {design.body_size_pt:.1f} pt, {design.columns} column(s)"
            ),
            payload={"grid": design.to_grid(), "measured": design.to_dict()},
            confidence=design.confidence,
            source=source,
        )
    ]


def _obligations(text: str, standard: str, source: str) -> list[dict[str, Any]]:
    """Delegated to foldok_learn, which already draws this line correctly."""
    try:
        from foldok_learn import extract
    except ImportError:  # pragma: no cover
        return []
    return [f.to_dict() for f in extract(text, source=source).findings]


def _obligation_proposals(obligations: list[dict[str, Any]], source: str) -> list[ShredProposal]:
    if not obligations:
        return []
    evidential = [o for o in obligations if o.get("artifact") in
                  ("measurement", "photo", "signature")]
    return [
        ShredProposal(
            kind="requirement",
            title=f"{len(obligations)} obligation(s), {len(evidential)} needing evidence",
            detail="citations only — no text from the standard was kept",
            payload={"obligations": obligations},
            confidence=0.7,
            source=source,
        )
    ]
